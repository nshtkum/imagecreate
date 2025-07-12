import streamlit as st
import pandas as pd
import requests
import time
import json
import re
import zipfile
import base64
from io import BytesIO
from docx import Document
from huggingface_hub import InferenceClient

# Configure page
st.set_page_config(page_title="SEO Content Automation", page_icon="📚", layout="wide")

# Display current AI provider info
if 'ai_provider' not in st.session_state:
    st.session_state['ai_provider'] = "Grok (X.AI)"

# Initialize session state
if "articles" not in st.session_state:
    st.session_state["articles"] = {}
if "images" not in st.session_state:
    st.session_state["images"] = {}
if "publish_log" not in st.session_state:
    st.session_state["publish_log"] = []

def init_hf_client():
    """Initialize Hugging Face client"""
    try:
        HF_TOKEN = st.secrets.get("HF_TOKEN")
        if not HF_TOKEN:
            st.warning("⚠️ HF_TOKEN not found in secrets. AI generation will be disabled.")
            return None
        return InferenceClient(
            model="stabilityai/stable-diffusion-3-medium",
            token=HF_TOKEN
        )
    except Exception as e:
        st.error(f"Error initializing HF client: {str(e)}")
        return None

def call_ai_for_metadata(keyword, intent, content_type, notes, api_key, provider):
    """Generate metadata using AI API"""
    prompt = f"""
You are a search volume estimator and content strategist for Indian students and professionals.

Given the keyword: "{keyword}"
Intent: {intent}
Content Type: {content_type}
Notes: {notes}

Tasks:
1. Estimate search volume as High, Medium, or Low based on Indian market
2. Create a high-CTR SEO title optimized for Indian audience (include year 2024/2025)
3. Generate a detailed outline with 5-7 bullet points including:
   - Introduction/What is section
   - Key features/benefits
   - Detailed information (eligibility, process, etc.)
   - Tables section (comparison/data)
   - FAQ section
   - Conclusion
4. Provide specific content instructions mentioning:
   - Target keyword usage frequency
   - Required tables (comparison, fees, eligibility, etc.)
   - FAQ requirements (5-8 questions)
   - Tone and style preferences
   - Indian context requirements

Respond in JSON format:
{{
  "volume": "High/Medium/Low",
  "seo_title": "Complete Guide to [Keyword] in India 2024 - Benefits, Process & FAQs",
  "outline": [
    "Introduction - What is [Keyword]?",
    "Key Features and Benefits of [Keyword]", 
    "Detailed [Keyword] Information and Requirements",
    "Comparison Table - [Keyword] Categories/Types",
    "Step-by-Step Application Process",
    "Eligibility Criteria and Documents Required",
    "Frequently Asked Questions (FAQs)",
    "Conclusion and Key Takeaways"
  ],
  "instructions": "Use keyword '[keyword]' 8-12 times naturally. Include 2 detailed tables: one for comparison/categories and one for fees/requirements. Add comprehensive FAQ section with 6-8 questions. Focus on Indian context with specific data. Use professional, informative tone. Include current 2024/2025 information."
}}
"""
    
    # Configure API based on provider
    if provider == "Grok (X.AI)":
        url = "https://api.x.ai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "grok-beta",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }
    elif provider == "Perplexity":
        url = "https://api.perplexity.ai/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "mistral-7b-instruct",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }
    elif provider == "OpenAI":
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "gpt-4o-mini",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }
    else:
        st.error(f"Unsupported provider: {provider}")
        return None

    try:
        response = requests.post(url, json=body, headers=headers)
        if response.status_code == 200:
            content = response.json()["choices"][0]["message"]["content"]
            # Try to parse JSON from response
            try:
                json_data = json.loads(content)
                return json_data
            except:
                # Fallback if JSON parsing fails
                return {
                    "volume": "Medium",
                    "seo_title": f"Complete Guide to {keyword} in India 2024 - Benefits, Eligibility & Process",
                    "outline": [
                        f"What is {keyword}? - Introduction and Overview",
                        f"Key Features and Benefits of {keyword}",
                        f"Eligibility Criteria for {keyword}",
                        f"Application Process and Required Documents",
                        f"Detailed Information Table - {keyword} Categories",
                        f"Frequently Asked Questions (FAQs) about {keyword}",
                        "Conclusion and Important Points"
                    ],
                    "instructions": f"Use '{keyword}' naturally 10-15 times. Include 2 tables: eligibility criteria and comparison table. Add FAQ section with 6-8 questions. Focus on Indian context with current data. Professional tone."
                }
        else:
            st.error(f"API Error ({provider}): {response.status_code} - {response.text}")
            return None
    except Exception as e:
        st.error(f"Request failed ({provider}): {str(e)}")
        return None

def generate_article(keyword, seo_title, outline, instructions, api_key, provider):
    """Generate complete article based on metadata"""
    outline_text = "\n".join([f"- {point}" for point in outline])
    
    prompt = f"""
You are an expert content writer specializing in educational content for Indian students and professionals.

Write a complete SEO-optimized article for:
Keyword: "{keyword}"
Title: "{seo_title}"

Structure using this outline:
{outline_text}

Special Instructions:
{instructions}

CRITICAL FORMATTING REQUIREMENTS:
1. Use the target keyword "{keyword}" naturally throughout the article (aim for 1-2% keyword density)
2. Include keyword variations and related terms
3. Structure with clear H1, H2, H3 headings using HTML tags
4. Include at least 1-2 detailed tables with relevant data
5. Add a comprehensive FAQ section at the end with 5-8 questions
6. Use bullet points and numbered lists where appropriate
7. Include specific data, statistics, and numbers
8. Write in HTML format with proper semantic tags

ARTICLE STRUCTURE TEMPLATE:
<h1>{seo_title}</h1>
<p>Introduction paragraph mentioning "{keyword}" and its importance...</p>

<h2>What is {keyword}?</h2>
<p>Detailed explanation...</p>

<h2>Key Features/Benefits of {keyword}</h2>
<ul>
<li>Feature 1 with explanation</li>
<li>Feature 2 with explanation</li>
</ul>

<h2>Detailed Information Table</h2>
<table border="1" cellpadding="8" cellspacing="0">
<tr><th>Parameter</th><th>Details</th></tr>
<tr><td>...</td><td>...</td></tr>
</table>

<h2>How to Apply/Process</h2>
<ol>
<li>Step 1</li>
<li>Step 2</li>
</ol>

<h2>Eligibility Criteria</h2>
<p>Detailed eligibility information...</p>

<h2>Frequently Asked Questions (FAQs)</h2>
<h3>Q1: What is {keyword}?</h3>
<p>A: Detailed answer...</p>

<h3>Q2: Who is eligible for {keyword}?</h3>
<p>A: Detailed answer...</p>

[Continue with 5-8 FAQs total]

<h2>Conclusion</h2>
<p>Summary mentioning "{keyword}" and key takeaways...</p>

Requirements:
- 1000-1500 words
- Natural keyword usage throughout
- Include specific Indian context and data
- Use tables for complex information
- Comprehensive FAQ section
- Professional, informative tone

Write the complete article now following this structure:
"""
    
    # Configure API based on provider
    if provider == "Grok (X.AI)":
        url = "https://api.x.ai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "grok-beta",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
            "max_tokens": 2000
        }
    elif provider == "Perplexity":
        url = "https://api.perplexity.ai/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "mistral-7b-instruct",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }
    elif provider == "OpenAI":
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "gpt-4o-mini",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
            "max_tokens": 2000
        }

    try:
        response = requests.post(url, json=body, headers=headers)
        if response.status_code == 200:
            article = response.json()["choices"][0]["message"]["content"]
            return article
        else:
            st.error(f"Article generation failed ({provider}): {response.status_code} - {response.text}")
            return None
    except Exception as e:
        st.error(f"Article generation error ({provider}): {str(e)}")
        return None

def apply_internal_links(article_content, anchor_map):
    """Apply internal links to article content"""
    linked_article = article_content
    
    # Replace first occurrence of each anchor (case-insensitive)
    for anchor, url in anchor_map.items():
        pattern = re.compile(rf"\b({re.escape(anchor)})\b", re.IGNORECASE)
        linked_article, n = pattern.subn(
            rf'<a href="{url}" target="_blank">\1</a>', 
            linked_article, 
            count=1
        )
    
    # Add related links table
    table_rows = "".join([
        f'<tr><td>{anchor}</td><td><a href="{url}" target="_blank">{url}</a></td></tr>'
        for anchor, url in anchor_map.items()
    ])
    
    link_table = f"""
<div style="margin-top: 30px;">
<h2>Related Links</h2>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse; width: 100%;">
<tr style="background-color: #f2f2f2;"><th>Topic</th><th>Link</th></tr>
{table_rows}
</table>
</div>
"""
    
    return linked_article + link_table

def generate_ai_image(prompt, hf_client):
    """Generate image using Hugging Face Inference Client"""
    if not hf_client:
        st.error("Hugging Face client not initialized")
        return None
    
    try:
        # Generate image using the inference client
        image = hf_client.text_to_image(prompt)
        
        # Convert PIL image to BytesIO
        img_buffer = BytesIO()
        image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        return img_buffer
    except Exception as e:
        st.error(f"Image generation error: {str(e)}")
        return None

def publish_to_wordpress(keyword, content, image_buffer, tags, wp_config, publish_now=True):
    """Publish article to WordPress"""
    wp_base = wp_config["base_url"]
    auth_str = f"{wp_config['username']}:{wp_config['password']}"
    auth_token = base64.b64encode(auth_str.encode()).decode("utf-8")
    headers = {"Authorization": f"Basic {auth_token}"}
    
    img_id = None
    
    # Upload image if provided
    if image_buffer:
        try:
            image_buffer.seek(0)
            img_data = image_buffer.read()
            img_headers = headers.copy()
            img_headers.update({
                "Content-Disposition": f"attachment; filename={keyword.replace(' ', '_')}.jpg",
                "Content-Type": "image/jpeg"
            })
            media_url = f"{wp_base}/wp-json/wp/v2/media"
            img_resp = requests.post(media_url, headers=img_headers, data=img_data)
            
            if img_resp.status_code == 201:
                img_id = img_resp.json()["id"]
            else:
                st.warning(f"Image upload failed for {keyword}: {img_resp.text}")
        except Exception as e:
            st.error(f"Image upload error: {str(e)}")
    
    # Create/get tags
    tag_ids = []
    if tags:
        for tag in [t.strip() for t in tags.split(",") if t.strip()]:
            try:
                # Check if tag exists
                tag_check = requests.get(f"{wp_base}/wp-json/wp/v2/tags?search={tag}", headers=headers)
                if tag_check.status_code == 200 and tag_check.json():
                    tag_ids.append(tag_check.json()[0]["id"])
                else:
                    # Create new tag
                    tag_create = requests.post(f"{wp_base}/wp-json/wp/v2/tags", headers=headers, json={"name": tag})
                    if tag_create.status_code == 201:
                        tag_ids.append(tag_create.json()["id"])
            except Exception as e:
                st.warning(f"Tag creation failed for '{tag}': {str(e)}")
    
    # Publish article
    post_data = {
        "title": keyword,
        "content": content,
        "status": "publish" if publish_now else "draft",
        "tags": tag_ids
    }
    
    if img_id:
        post_data["featured_media"] = img_id
    
    try:
        post_resp = requests.post(f"{wp_base}/wp-json/wp/v2/posts", headers=headers, json=post_data)
        if post_resp.status_code == 201:
            post_url = post_resp.json()["link"]
            return {"success": True, "url": post_url}
        else:
            return {"success": False, "error": post_resp.text}
    except Exception as e:
        return {"success": False, "error": str(e)}

# Main App Interface
st.title("📚 SEO Content Automation Pipeline")
st.markdown("Upload topics → Generate metadata → Create articles → Add images → Publish to WordPress")

# Show current provider
if api_key:
    st.success(f"✅ Connected to {ai_provider}")
else:
    st.info(f"📡 Ready to use {ai_provider} - Please add API key in sidebar")

# Sidebar for API configuration
st.sidebar.header("🔧 API Configuration")

# AI Model Selection
ai_provider = st.sidebar.selectbox(
    "Choose AI Provider",
    ["Grok (X.AI)", "OpenAI"],
    help="Select your preferred AI provider for content generation"
)

# Check API keys
api_key = None
if ai_provider == "Grok (X.AI)":
    api_key = st.secrets.get("GROK_API_KEY")
elif ai_provider == "OpenAI":
    api_key = st.secrets.get("OPENAI_API_KEY")

# Initialize HF client
hf_client = init_hf_client()

# WordPress Config
st.sidebar.header("🌐 WordPress Settings")
wp_base_url = st.secrets.get("WP_BASE_URL", "")
wp_username = st.secrets.get("WP_USERNAME", "")
wp_password = st.secrets.get("WP_PASSWORD", "")

if wp_base_url and wp_username and wp_password:
    st.sidebar.success("✅ WordPress configured")
else:
    st.sidebar.warning("⚠️ WordPress not configured in secrets")

# Tab interface
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "📋 Topic Upload", 
    "📝 Article Generation", 
    "🔗 Internal Links", 
    "🖼️ Images", 
    "🚀 WordPress Publish", 
    "📊 Export & Logs"
])

with tab1:
    st.header("📋 Step 1: Topic Metadata Generation")
    
    uploaded_file = st.file_uploader("Upload Excel with keywords", type=["xlsx"])
    
    if uploaded_file and api_key:
        df = pd.read_excel(uploaded_file)
        
        required_cols = ["Keyword", "Intent", "Content Type", "Notes"]
        if not all(col in df.columns for col in required_cols):
            st.error(f"❌ Your Excel must contain columns: {', '.join(required_cols)}")
        else:
            st.success(f"✅ Found {len(df)} topics to process")
            
            if st.button("🚀 Generate Metadata for All Topics"):
                progress_bar = st.progress(0)
                results = []
                
                for i, row in df.iterrows():
                    st.info(f"Processing: {row['Keyword']} ({i+1}/{len(df)})")
                    
                    output = call_ai_for_metadata(
                        row["Keyword"],
                        row["Intent"],
                        row["Content Type"],
                        row["Notes"],
                        api_key,
                        ai_provider
                    )
                    
                    if output:
                        results.append({
                            "Keyword": row["Keyword"],
                            "Intent": row["Intent"],
                            "Content Type": row["Content Type"],
                            "Volume": output["volume"],
                            "SEO Title": output["seo_title"],
                            "Outline": "\n".join(output["outline"]),
                            "Instructions": output["instructions"]
                        })
                    
                    progress_bar.progress((i + 1) / len(df))
                    time.sleep(1.5)  # API rate limiting
                
                if results:
                    st.session_state["metadata_df"] = pd.DataFrame(results)
                    st.success("✅ Metadata generated successfully!")
                    st.dataframe(st.session_state["metadata_df"], use_container_width=True)
                    
                    # Download option
                    csv = st.session_state["metadata_df"].to_csv(index=False)
                    st.download_button(
                        "⬇️ Download Metadata CSV",
                        csv,
                        file_name="topic_metadata.csv",
                        mime="text/csv"
                    )
    
    elif not api_key:
        st.warning(f"⚠️ Please enter your {ai_provider} API key in the sidebar")

with tab2:
    st.header("📝 Step 2: Article Generation")
    
    if "metadata_df" in st.session_state and not st.session_state["metadata_df"].empty:
        metadata_df = st.session_state["metadata_df"]
        
        # Single article generation
        st.subheader("Generate Single Article")
        selected_keyword = st.selectbox("Select a topic", metadata_df["Keyword"])
        
        if selected_keyword:
            selected_row = metadata_df[metadata_df["Keyword"] == selected_keyword].iloc[0]
            
            # Allow editing of metadata
            seo_title = st.text_input("SEO Title", selected_row["SEO Title"])
            outline = st.text_area("Outline", selected_row["Outline"])
            instructions = st.text_area("Instructions", selected_row["Instructions"])
            
            if st.button("Generate Article") and api_key:
                with st.spinner("Generating article..."):
                    article = generate_article(
                        selected_keyword,
                        seo_title,
                        outline.split("\n"),
                        instructions,
                        api_key,
                        ai_provider
                    )
                    
                    if article:
                        st.session_state["articles"][selected_keyword] = article
                        st.success("✅ Article generated!")
                        st.markdown("### Preview:")
                        st.markdown(article, unsafe_allow_html=True)
        
        # Article structure preview
        if st.session_state["articles"]:
            st.subheader("📋 Article Structure Preview")
            with st.expander("Click to see expected article structure"):
                st.markdown("""
                **Your articles will include:**
                
                1. **SEO-Optimized Title** (H1 with target keyword)
                2. **Introduction** - What is [Keyword]?
                3. **Key Features/Benefits** (with bullet points)
                4. **Detailed Information Tables** (eligibility, comparison, fees)
                5. **Step-by-Step Process** (numbered lists)
                6. **Eligibility Criteria** (with specific requirements)
                7. **Comprehensive FAQ Section** (6-8 questions)
                8. **Conclusion** (summary with key takeaways)
                
                **SEO Features:**
                - Target keyword used 10-15 times naturally
                - Structured headings (H1, H2, H3)
                - Rich tables with relevant data
                - FAQ section for long-tail keywords
                - 1000-1500 words length
                - Indian context and current data
                """)
        
        # Bulk generation
        st.subheader("Generate All Articles")
        if st.button("🚀 Generate All Articles") and api_key:
            progress_bar = st.progress(0)
            
            for i, row in metadata_df.iterrows():
                keyword = row["Keyword"]
                st.info(f"Generating article for: {keyword}")
                
                article = generate_article(
                    keyword,
                    row["SEO Title"],
                    row["Outline"].split("\n"),
                    row["Instructions"],
                    api_key,
                    ai_provider
                )
                
                if article:
                    st.session_state["articles"][keyword] = article
                
                progress_bar.progress((i + 1) / len(metadata_df))
                time.sleep(1.5)
            
            st.success(f"✅ Generated {len(st.session_state['articles'])} articles!")
    
    else:
        st.info("⚠️ Please complete Step 1 first to generate metadata")

with tab3:
    st.header("🔗 Step 3: Internal Linking")
    
    if st.session_state["articles"]:
        st.subheader("Upload Anchor Links Mapping")
        mapping_file = st.file_uploader(
            "Upload CSV with 'anchor,url' columns or JSON file",
            type=["csv", "json"]
        )
        
        if mapping_file:
            try:
                if mapping_file.name.endswith(".csv"):
                    link_df = pd.read_csv(mapping_file)
                    if "anchor" in link_df.columns and "url" in link_df.columns:
                        anchor_map = dict(zip(link_df["anchor"], link_df["url"]))
                        st.success(f"✅ Loaded {len(anchor_map)} anchor links")
                        st.dataframe(link_df)
                    else:
                        st.error("❌ CSV must have 'anchor' and 'url' columns")
                        anchor_map = {}
                else:  # JSON
                    anchor_map = json.load(mapping_file)
                    st.success(f"✅ Loaded {len(anchor_map)} anchor links")
                    st.json(anchor_map)
                
                if anchor_map:
                    selected_articles = st.multiselect(
                        "Select articles to apply internal links",
                        options=list(st.session_state["articles"].keys()),
                        default=list(st.session_state["articles"].keys())
                    )
                    
                    if st.button("🔗 Apply Internal Links"):
                        for keyword in selected_articles:
                            original_article = st.session_state["articles"][keyword]
                            linked_article = apply_internal_links(original_article, anchor_map)
                            st.session_state["articles"][keyword] = linked_article
                        
                        st.success("✅ Internal links applied successfully!")
                        
            except Exception as e:
                st.error(f"Error processing mapping file: {str(e)}")
    else:
        st.info("⚠️ Please generate articles first")

with tab4:
    st.header("🖼️ Step 4: Featured Images")
    
    if st.session_state["articles"]:
        keywords = list(st.session_state["articles"].keys())
        
        # Manual image upload
        st.subheader("Upload Images Manually")
        for keyword in keywords:
            with st.expander(f"📌 {keyword}"):
                uploaded_img = st.file_uploader(
                    f"Upload image for '{keyword}'",
                    type=["jpg", "png", "jpeg"],
                    key=f"manual_img_{keyword}"
                )
                if uploaded_img:
                    st.session_state["images"][keyword] = BytesIO(uploaded_img.read())
                    st.image(uploaded_img, caption="Uploaded", width=300)
        
        # AI image generation
        st.subheader("Generate AI Images")
        if hf_client:
            selected_for_ai = st.selectbox("Select topic for AI image", keywords)
            
            if selected_for_ai:
                default_prompt = f"High-quality professional illustration representing {selected_for_ai}, educational content, clean design"
                image_prompt = st.text_input("Image generation prompt", default_prompt)
                
                if st.button("🎨 Generate AI Image"):
                    with st.spinner("Generating image..."):
                        image_buffer = generate_ai_image(image_prompt, hf_client)
                        
                        if image_buffer:
                            st.session_state["images"][selected_for_ai] = image_buffer
                            st.success("✅ AI image generated!")
                            st.image(image_buffer, caption="AI Generated", width=300)
        else:
            st.info("⚠️ Hugging Face client not available for AI image generation")
    
    else:
        st.info("⚠️ Please generate articles first")

with tab5:
    st.header("🚀 Step 5: WordPress Publishing")
    
    if st.session_state["articles"] and all([wp_base_url, wp_username, wp_password]):
        wp_config = {
            "base_url": wp_base_url,
            "username": wp_username,
            "password": wp_password
        }
        
        # Publishing options
        st.subheader("Publishing Settings")
        selected_to_publish = st.multiselect(
            "Select articles to publish",
            options=list(st.session_state["articles"].keys()),
            default=list(st.session_state["articles"].keys())
        )
        
        tags_input = st.text_input("Tags (comma-separated)", "Education, India, Guide")
        publish_immediately = st.checkbox("Publish immediately", value=True)
        
        if st.button("🚀 Publish to WordPress"):
            if not selected_to_publish:
                st.error("Please select at least one article to publish")
            else:
                progress_bar = st.progress(0)
                published_count = 0
                
                for i, keyword in enumerate(selected_to_publish):
                    st.info(f"Publishing: {keyword}")
                    
                    content = st.session_state["articles"][keyword]
                    image_buffer = st.session_state["images"].get(keyword)
                    
                    result = publish_to_wordpress(
                        keyword,
                        content,
                        image_buffer,
                        tags_input,
                        wp_config,
                        publish_immediately
                    )
                    
                    log_entry = {
                        "keyword": keyword,
                        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
                        "status": "Success" if result["success"] else "Failed",
                        "url": result.get("url", ""),
                        "error": result.get("error", "")
                    }
                    
                    st.session_state["publish_log"].append(log_entry)
                    
                    if result["success"]:
                        published_count += 1
                        st.success(f"✅ Published: [{keyword}]({result['url']})")
                    else:
                        st.error(f"❌ Failed to publish '{keyword}': {result['error']}")
                    
                    progress_bar.progress((i + 1) / len(selected_to_publish))
                    time.sleep(1)  # Rate limiting
                
                st.success(f"🎉 Publishing complete! {published_count}/{len(selected_to_publish)} articles published successfully.")
    
    elif not st.session_state["articles"]:
        st.info("⚠️ Please generate articles first")
    else:
        st.info("⚠️ Please configure WordPress settings in the sidebar")

with tab6:
    st.header("📊 Export & Logs")
    
    # Publishing logs
    if st.session_state["publish_log"]:
        st.subheader("Publishing Logs")
        log_df = pd.DataFrame(st.session_state["publish_log"])
        st.dataframe(log_df, use_container_width=True)
        
        # Download logs
        csv_log = log_df.to_csv(index=False)
        st.download_button(
            "📥 Download Publish Log",
            csv_log,
            file_name="publish_log.csv",
            mime="text/csv"
        )
    
    # Export all content
    if st.session_state["articles"]:
        st.subheader("Export All Content")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📦 Download All Articles (HTML)"):
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                    for keyword, content in st.session_state["articles"].items():
                        filename = f"{keyword.replace(' ', '_')}.html"
                        zip_file.writestr(filename, content)
                
                zip_buffer.seek(0)
                st.download_button(
                    "⬇️ Download HTML Files",
                    zip_buffer,
                    file_name="all_articles_html.zip",
                    mime="application/zip"
                )
        
        with col2:
            if st.button("📄 Download All Articles (Word)"):
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                    for keyword, content in st.session_state["articles"].items():
                        # Create Word document
                        doc = Document()
                        doc.add_heading(keyword, 0)
                        
                        # Simple HTML to text conversion for Word
                        clean_content = re.sub('<[^<]+?>', '', content)
                        paragraphs = clean_content.split('\n')
                        
                        for para in paragraphs:
                            if para.strip():
                                doc.add_paragraph(para.strip())
                        
                        # Save to buffer
                        word_buffer = BytesIO()
                        doc.save(word_buffer)
                        word_buffer.seek(0)
                        
                        filename = f"{keyword.replace(' ', '_')}.docx"
                        zip_file.writestr(filename, word_buffer.read())
                
                zip_buffer.seek(0)
                st.download_button(
                    "⬇️ Download Word Files",
                    zip_buffer,
                    file_name="all_articles_word.zip",
                    mime="application/zip"
                )
    
    else:
        st.info("⚠️ No content available for export")

# Footer
st.markdown("---")
st.markdown("**SEO Content Automation Pipeline** - Generate, optimize, and publish content at scale")
